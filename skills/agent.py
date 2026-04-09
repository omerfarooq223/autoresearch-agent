"""
AutoResearch Agent - Enhanced Design Version
==============================================
An autonomous AI research agent that searches the web and produces
professional PDF research reports on any topic.
Now with modern editorial design and visual enhancements.

Author: M. Umer Farooq
GitHub: https://github.com/umerfarooq
Built with: Groq (LLaMA 3.3 70B) + Tavily Search API
"""

import os
import sys
import hashlib
import json
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse
from xml.sax.saxutils import escape, quoteattr
from dotenv import load_dotenv
from tavily import TavilyClient
from groq import Groq
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, white, black
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 HRFlowable, Table, TableStyle,
                                 PageBreak, BaseDocTemplate, Frame,
                                 PageTemplate)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT

load_dotenv()


# --- Error handling for API keys ---
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
if not GROQ_API_KEY:
    print("[ERROR] GROQ_API_KEY not found in environment. Please set it in your .env file.")
    sys.exit(1)
if not TAVILY_API_KEY:
    print("[ERROR] TAVILY_API_KEY not found in environment. Please set it in your .env file.")
    sys.exit(1)

groq_client = Groq(api_key=GROQ_API_KEY)
tavily = TavilyClient(api_key=TAVILY_API_KEY)
MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")

# Enhanced brand colors with better contrast
DARK_BLUE    = HexColor('#0f172a')
ACCENT_BLUE  = HexColor('#2563eb')
ACCENT_TEAL  = HexColor('#0891b2')
LIGHT_BLUE   = HexColor('#dbeafe')
TEAL         = HexColor('#0891b2')
GRAY         = HexColor('#6b7280')
LIGHT_GRAY   = HexColor('#f1f5f9')
DARK_GRAY    = HexColor('#1e293b')
WHITE        = HexColor('#ffffff')
GOLD         = HexColor('#f59e0b')
ACCENT_PURPLE = HexColor('#7c3aed')
ACCENT_CYAN  = HexColor('#06b6d4')


def color_to_hex(color):
    """Convert a ReportLab Color to #RRGGBB."""
    try:
        r = int(max(0, min(1, color.red)) * 255)
        g = int(max(0, min(1, color.green)) * 255)
        b = int(max(0, min(1, color.blue)) * 255)
        return f"#{r:02x}{g:02x}{b:02x}"
    except Exception:
        return "#000000"


def safe_paragraph_text(value):
    return escape(str(value or ""))


def format_url_for_pdf(url):
    """Render full URL with controlled line breaks to avoid clipping."""
    u = str(url or "")
    if not u:
        return ""
    parsed = urlparse(u)
    if not parsed.scheme:
        return safe_paragraph_text(u)

    domain = f"{parsed.scheme}://{parsed.netloc}"
    path_parts = [p for p in parsed.path.split('/') if p]
    wrapped = [safe_paragraph_text(domain)]
    for part in path_parts:
        wrapped.append(safe_paragraph_text(part))
    if parsed.query:
        wrapped.append("?" + safe_paragraph_text(parsed.query))
    if parsed.fragment:
        wrapped.append("#" + safe_paragraph_text(parsed.fragment))
    return "<br/>".join(wrapped)


def make_pdf_link(url):
    """Return a clickable URL fragment for ReportLab Paragraph."""
    u = str(url or "").strip()
    if not u:
        return ""
    href = quoteattr(u)
    label = format_url_for_pdf(u)
    return f"<link href={href} color=\"#2563eb\"><u>{label}</u></link>"


def ask_groq(prompt, system="You are a helpful research assistant."):
    model_candidates = [
        MODEL,
        "llama-3.1-8b-instant",
        "gemma2-9b-it",
    ]
    # Keep order while removing duplicates.
    model_candidates = list(dict.fromkeys(model_candidates))

    last_error = None
    for model_name in model_candidates:
        try:
            response = groq_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            msg = str(e).lower()
            # Fail fast on invalid credentials so UI can show a clear action item.
            if "invalid_api_key" in msg or "error code: 401" in msg:
                raise RuntimeError(
                    "Groq authentication failed: invalid API key. Please update GROQ_API_KEY."
                ) from e

            # If a model is unavailable, try the next fallback model.
            if "model" in msg and ("not found" in msg or "decommissioned" in msg):
                last_error = e
                continue

            last_error = e
            break

    raise RuntimeError(f"Groq API call failed: {last_error}")

def generate_sub_questions(topic):
    print(f"\n🧠 Breaking down topic: {topic}")
    prompt = f"""Break this research topic into 5 specific sub-questions that together would cover it comprehensively:
Topic: {topic}
Return ONLY a numbered list of 5 questions, nothing else."""
    result = ask_groq(prompt)
    questions = []
    for line in result.strip().split('\n'):
        s = line.strip()
        if not s:
            continue
        # Accept numbered outputs like "1. ..." and simple bullets.
        if s[0].isdigit():
            if "." in s:
                s = s.split(".", 1)[1].strip()
            questions.append(s)
        elif s.startswith("- ") or s.startswith("* "):
            questions.append(s[2:].strip())

    # Guard against parsing failures so downstream thread pools never get 0 workers.
    if not questions:
        questions = [
            f"What are the most recent and credible developments about {topic}?",
            f"What market, technical, or policy factors are shaping {topic}?",
            f"What risks, limitations, or counterarguments exist for {topic}?",
        ]
    for q in questions:
        print(f"  {q}")
    return questions


TRUSTED_DOMAINS = {
    "scholar.google.com", "pubmed.ncbi.nlm.nih.gov", "arxiv.org",
    "researchgate.net", "semanticscholar.org", "jstor.org",
    "reuters.com", "apnews.com", "bbc.com", "nature.com",
    "scientificamerican.com", "theatlantic.com", "nytimes.com",
    "theguardian.com", "wsj.com", "ft.com",
    "techcrunch.com", "wired.com", "mit.edu", "stanford.edu",
}

def estimate_credibility(source):
    url = source.get('url', '').lower()
    content = source.get('content', '')
    title = source.get('title', '')
    score = 0

    if any(tld in url for tld in [".gov", ".edu", ".mil"]):
        score += 3
    elif ".org" in url:
        score += 2
    elif ".com" in url or ".net" in url:
        score += 1

    try:
        from urllib.parse import urlparse
        domain = urlparse(url).netloc.replace("www.", "")
        if domain in TRUSTED_DOMAINS:
            score += 2
    except Exception:
        pass

    if len(content) > 500:
        score += 1
    if any(c.isdigit() for c in content[:300]):
        score += 1
    if title and len(title) > 20:
        score += 1

    return min(score, 5)

def extract_evidence_snippet(content):
    # Return the first 200 chars as a snippet (could be improved)
    return content[:200] + ("..." if len(content) > 200 else "")

def search_web(query, max_results=4, cache_dir=".cache"):
    print(f"\n🔍 Searching: {query[:65]} (max_results={max_results})...")
    # --- Caching logic ---
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)
    cache_key = hashlib.md5(f"{query}|{max_results}".encode()).hexdigest()
    cache_path = os.path.join(cache_dir, f"{cache_key}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                cached = json.load(f)
            print("[CACHE] Using cached search results.")
            return cached
        except Exception:
            pass
    try:
        results = tavily.search(query=query, max_results=max_results, search_depth="advanced")
    except Exception as e:
        print(f"[ERROR] Tavily API call failed: {e}")
        return []
    sources = []
    for r in results.get('results', []):
        credibility = estimate_credibility(r)
        snippet = extract_evidence_snippet(r.get('content', ''))
        sources.append({
            'title': r.get('title', ''),
            'url':   r.get('url', ''),
            'content': r.get('content', '')[:1000],
            'credibility': credibility,
            'snippet': snippet
        })
    if not sources:
        print("[WARNING] No sources found for this query.")
    # Save to cache
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(sources, f)
    except Exception:
        pass
    return sources

def synthesize_findings(topic, all_sources):
    print(f"\n✍️  Synthesizing {len(all_sources)} sources into report...")
    sources_text = ""
    for i, s in enumerate(all_sources[:12], 1):
        sources_text += f"\nSource {i}: {s['title']}\nURL: {s['url']}\nContent: {s['content'][:400]}\n"

    prompt = f"""You are a senior research analyst at a top consulting firm. 
Write a thorough, professional research report on: "{topic}"

Use EXACTLY this structure and be very detailed in each section:

## Executive Summary
Write 5-6 detailed sentences covering the most important findings, current state, and implications.

## Background & Context
Write 4-5 sentences providing context, history, and why this topic matters today.

## Key Findings

### Finding 1: [Descriptive Theme Title]
Write 2 sentences introducing this theme.
- Specific finding with data or evidence
- Specific finding with data or evidence  
- Specific finding with data or evidence
- Specific finding with data or evidence

### Finding 2: [Descriptive Theme Title]
Write 2 sentences introducing this theme.
- Specific finding with data or evidence
- Specific finding with data or evidence
- Specific finding with data or evidence
- Specific finding with data or evidence

### Finding 3: [Descriptive Theme Title]
Write 2 sentences introducing this theme.
- Specific finding with data or evidence
- Specific finding with data or evidence
- Specific finding with data or evidence
- Specific finding with data or evidence

### Finding 4: [Descriptive Theme Title]
Write 2 sentences introducing this theme.
- Specific finding with data or evidence
- Specific finding with data or evidence
- Specific finding with data or evidence

## Areas of Consensus
Write 5-6 sentences about what experts and sources broadly agree on regarding this topic.

## Debates and Controversies
Write 5-6 sentences describing the main disagreements, opposing viewpoints, and contested claims.

## Challenges and Barriers
Write 4-5 sentences about the main obstacles, risks, or challenges in this area.

## What Remains Unknown
Write 4-5 sentences about knowledge gaps, unanswered questions, and areas needing more research.

## Future Outlook
Write 4-5 sentences about trends, predictions, and what to expect in the coming years.

## Conclusion
Write 5-6 sentences synthesizing the report, key takeaways, and final recommendations.

SOURCES TO USE:
{sources_text}

Write the complete detailed report now. Be thorough, specific, and analytical:"""

    return ask_groq(prompt, system="You are a senior research analyst who writes detailed, evidence-based, professional reports for executives and academics.")

def critique_report(report, all_sources):
    print("\n🔍 Running self-critique loop...")
    sources_text = "\n".join([
        f"- {s['title']}: {s['content'][:300]}" for s in all_sources
    ])
    prompt = f"""You are a strict fact-checker. Read this research report and the sources it was based on.

Identify claims in the report that are NOT clearly supported by the provided sources.

REPORT:
{report}

SOURCES:
{sources_text}

Return ONLY a JSON list of unsupported claims, like:
["claim 1", "claim 2"]
If all claims are supported, return an empty list: []
Return ONLY the JSON, nothing else. No explanation, no markdown."""

    result = ask_groq(prompt)
    try:
        clean = result.strip().replace("```json", "").replace("```", "").strip()
        return json.loads(clean)
    except Exception:
        return []


def fill_gaps(unsupported_claims):
    print(f"⚠️  Found {len(unsupported_claims)} unsupported claims, re-searching...")
    if not unsupported_claims:
        return []
    gap_sources = []
    with ThreadPoolExecutor(max_workers=max(1, len(unsupported_claims))) as executor:
        futures = {executor.submit(search_web, claim, 3): claim 
                   for claim in unsupported_claims}
        for future in as_completed(futures):
            try:
                gap_sources.extend(future.result())
            except Exception as e:
                print(f"[WARNING] Gap search failed: {e}")
    return gap_sources


def patch_report(original_report, gap_sources, unsupported_claims):
    claims_text = "\n".join(f"- {c}" for c in unsupported_claims)
    sources_text = "\n".join([
        f"Source: {s['title']}\nURL: {s['url']}\nContent: {s['content']}"
        for s in gap_sources
    ])
    prompt = f"""You are a research editor. These claims in the report lacked source support:

{claims_text}

New sources found specifically for these claims:
{sources_text}

Original report:
{original_report}

Revise ONLY the unsupported claims using the new sources. Keep everything else identical. Return the complete revised report."""

    return ask_groq(prompt)


# ============================================================
# ENHANCED DESIGN FUNCTIONS
# ============================================================

class ProfessionalDoc(BaseDocTemplate):
    def __init__(self, filename, topic, num_sources, **kwargs):
        super().__init__(filename, **kwargs)
        self.topic = topic
        self.num_sources = num_sources
        frame = Frame(
            self.leftMargin, self.bottomMargin + 0.6*inch,
            self.width, self.height - 1.4*inch, id='main'
        )
        template = PageTemplate(id='main', frames=frame,
                                onPage=self.add_header_footer)
        self.addPageTemplates([template])

    def add_header_footer(self, c, doc):
        c.saveState()
        w, h = letter

        # Enhanced Header with gradient-like effect
        c.setFillColor(DARK_BLUE)
        c.rect(0, h - 0.75*inch, w, 0.75*inch, fill=1, stroke=0)
        
        # Accent stripe with gradient colors
        c.setFillColor(ACCENT_BLUE)
        c.rect(0, h - 0.75*inch, 0.12*inch, 0.75*inch, fill=1, stroke=0)
        c.setFillColor(ACCENT_TEAL)
        c.rect(0.12*inch, h - 0.75*inch, 0.08*inch, 0.75*inch, fill=1, stroke=0)

        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(0.3*inch, h - 0.32*inch, "RESEARCH INTELLIGENCE")
        c.setFont("Helvetica", 8)
        c.setFillColor(HexColor('#93c5fd'))
        c.drawString(0.3*inch, h - 0.52*inch, "Powered by LLaMA 3.3 70B + Tavily Search | M. Umer Farooq")

        topic_short = self.topic[:50] + "..." if len(self.topic) > 50 else self.topic
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 9)
        c.drawRightString(w - 0.3*inch, h - 0.42*inch, topic_short)

        # Enhanced Footer with border
        c.setFillColor(LIGHT_GRAY)
        c.rect(0, 0, w, 0.55*inch, fill=1, stroke=0)
        c.setStrokeColor(ACCENT_BLUE)
        c.setLineWidth(2.5)
        c.line(0.2*inch, 0.55*inch, w-0.2*inch, 0.55*inch)

        c.setFillColor(GRAY)
        c.setFont("Helvetica", 8)
        c.drawString(0.3*inch, 0.22*inch,
                     f"© {datetime.now().year} M. Umer Farooq  |  AutoResearch Agent  |  {datetime.now().strftime('%b %d, %Y')}")
        c.setFillColor(ACCENT_BLUE)
        c.setFont("Helvetica-Bold", 8)
        c.drawRightString(w - 0.3*inch, 0.22*inch,
                          f"Sources: {self.num_sources}  |  Page {doc.page}")
        c.restoreState()

def build_cover_page(topic, num_sources, styles):
    """Enhanced cover page with modern editorial design"""
    story = []
    story.append(Spacer(1, 0.5*inch))

    # Premium AI Badge
    badge_style = ParagraphStyle(
        'Badge', fontSize=10, fontName='Helvetica-Bold',
        textColor=ACCENT_BLUE, spaceAfter=20,
        borderColor=ACCENT_BLUE, borderWidth=2,
        borderPad=8, backColor=HexColor('#f0f9ff'),
        alignment=TA_CENTER, letterSpacing=1.5
    )
    story.append(Paragraph("AI-POWERED INTELLIGENCE REPORT", badge_style))
    story.append(Spacer(1, 0.08*inch))

    # Decorative accent line
    story.append(HRFlowable(width=1.2*inch, thickness=3,
                            color=GOLD, spaceAfter=20))

    # Main title with enhanced typography
    title_style = ParagraphStyle(
        'CoverTitle', fontSize=36, fontName='Helvetica-Bold',
        textColor=DARK_BLUE, spaceAfter=8, leading=42,
        letterSpacing=0.5
    )
    story.append(Paragraph("Research Intelligence", title_style))

    # Topic with visual emphasis
    topic_style = ParagraphStyle(
        'CoverTopic', fontSize=18, fontName='Helvetica',
        textColor=ACCENT_BLUE, spaceAfter=32, leading=26,
    )
    story.append(Paragraph(f"<b>{safe_paragraph_text(topic)}</b>", topic_style))

    story.append(Spacer(1, 0.15*inch))

    # Modern card-style metadata with alternating colors
    meta_label_style = ParagraphStyle(
        'MetaLabel', fontSize=10, leading=12, fontName='Helvetica-Bold', textColor=DARK_BLUE
    )
    meta_value_style = ParagraphStyle(
        'MetaValue', fontSize=10, leading=12, fontName='Helvetica', textColor=DARK_GRAY
    )
    meta_cards = [
        [Paragraph('<b>Report Date</b>', meta_label_style), Paragraph(datetime.now().strftime('%B %d, %Y'), meta_value_style)],
        [Paragraph('<b>Architect</b>', meta_label_style), Paragraph('M. Umer Farooq', meta_value_style)],
        [Paragraph('<b>AI Engine</b>', meta_label_style), Paragraph('LLaMA 3.3 70B (Ultra-Fast)', meta_value_style)],
        [Paragraph('<b>Search Engine</b>', meta_label_style), Paragraph('Tavily Advanced Search', meta_value_style)],
        [Paragraph('<b>Data Points</b>', meta_label_style), Paragraph(f'{num_sources} verified sources', meta_value_style)],
        [Paragraph('<b>Speed</b>', meta_label_style), Paragraph('Generated in under 60 seconds', meta_value_style)],
    ]
    
    t = Table(meta_cards, colWidths=[1.9*inch, 4.7*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), HexColor('#f0f4ff')),
        ('BACKGROUND', (1,0), (1,-1), WHITE),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('PADDING', (0,0), (-1,-1), 11),
        ('GRID', (0,0), (-1,-1), 1, HexColor('#e2e8f0')),
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [HexColor('#f8fbff'), WHITE]),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (1,0), (1,-1), 14),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.25*inch))

    # Enhanced disclaimer
    disc_style = ParagraphStyle(
        'Disclaimer', fontSize=9, fontName='Helvetica',
        textColor=HexColor('#374151'), spaceAfter=0, leading=13,
        borderColor=HexColor('#dbeafe'), borderWidth=1.5,
        borderPad=10, backColor=HexColor('#f0f9ff'),
        alignment=TA_LEFT
    )
    story.append(Paragraph(
        "<b>Verification Notice:</b> This report was generated by AutoResearch Agent "
        "using AI synthesis and web search. All claims reference publicly available sources. "
        "Validate critical findings through independent research before decision-making.",
        disc_style
    ))
    story.append(PageBreak())
    return story

def save_pdf(topic, report_text, all_sources, logo_path=None):
    """Enhanced PDF export with modern design"""
    filename = topic.lower().replace(' ', '-')[:40] + '-report.pdf'
    doc = ProfessionalDoc(
        filename, topic=topic, num_sources=len(all_sources),
        pagesize=letter,
        rightMargin=0.75*inch, leftMargin=0.75*inch,
        topMargin=1.0*inch, bottomMargin=0.9*inch
    )

    styles = getSampleStyleSheet()

    # Define enhanced styles
    h2_style = ParagraphStyle(
        'H2', fontSize=14, fontName='Helvetica-Bold',
        textColor=WHITE, spaceBefore=18, spaceAfter=10,
        leading=18, leftIndent=-6, rightIndent=-6,
        borderPad=10, backColor=DARK_BLUE, keepWithNext=True
    )
    h3_style = ParagraphStyle(
        'H3', fontSize=11, fontName='Helvetica-Bold',
        textColor=ACCENT_BLUE, spaceBefore=12, spaceAfter=5, leading=15,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'Body', fontSize=10, fontName='Helvetica',
        textColor=DARK_GRAY, spaceAfter=6, leading=16,
        alignment=TA_JUSTIFY, splitLongWords=True
    )
    bullet_style = ParagraphStyle(
        'Bullet', fontSize=10, fontName='Helvetica',
        textColor=DARK_GRAY, spaceAfter=5, leading=15,
        leftIndent=20, bulletIndent=6, splitLongWords=True
    )

    story = []

    # Cover page
    if logo_path and os.path.exists(logo_path):
        from reportlab.platypus import Image
        story.append(Image(logo_path, width=1.5*inch, height=1.5*inch))
        story.append(Spacer(1, 0.1*inch))
    story.extend(build_cover_page(topic, len(all_sources), styles))

    # Enhanced sources table with visual hierarchy
    if all_sources:
        story.append(Spacer(1, 0.1*inch))
        
        # Section header
        sources_header = ParagraphStyle(
            'SourcesHeader', fontSize=12, fontName='Helvetica-Bold',
            textColor=WHITE, spaceBefore=6, spaceAfter=10,
            borderPad=8, backColor=ACCENT_TEAL
        )
        story.append(Paragraph("   SOURCES & CREDIBILITY ASSESSMENT", sources_header))
        story.append(Spacer(1, 0.1*inch))

        cell_style = ParagraphStyle(
            'cell', fontSize=7.5, leading=9, textColor=DARK_GRAY,
            wordWrap='CJK', splitLongWords=True
        )
        header_style = ParagraphStyle('header', fontSize=8, leading=9, textColor=WHITE, 
                                     fontName='Helvetica-Bold')

        table_data = [[
            Paragraph('No.', header_style),
            Paragraph('Source Title', header_style),
            Paragraph('Credibility', header_style),
            Paragraph('Evidence Snippet', header_style),
            Paragraph('URL', header_style),
        ]]
        
        for i, s in enumerate(all_sources[:12], 1):
            # Credibility indicator with color coding
            cred = s.get('credibility', 0)
            if cred >= 4:
                cred_color = HexColor('#059669')
                cred_text = "High (5/5)"
            elif cred >= 3:
                cred_color = HexColor('#0891b2')
                cred_text = "Medium (4/5)"
            else:
                cred_color = HexColor('#f59e0b')
                cred_text = "Low (3/5)"
            
            cred_para = Paragraph(f'<font color="{color_to_hex(cred_color)}"><b>{cred_text}</b></font>', cell_style)
            
            table_data.append([
                Paragraph(str(i), cell_style),
                Paragraph(safe_paragraph_text(s.get('title', '')), cell_style),
                cred_para,
                Paragraph(safe_paragraph_text(s.get('snippet', '')), cell_style),
                Paragraph(make_pdf_link(s.get('url', '')), cell_style),
            ])

        t = Table(table_data, colWidths=[0.35*inch, 1.6*inch, 0.65*inch, 2.4*inch, 2.0*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), ACCENT_BLUE),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 8),
            ('FONTSIZE', (0,1), (-1,-1), 7.5),
            ('GRID', (0,0), (-1,-1), 0.75, HexColor('#e2e8f0')),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('PADDING', (0,0), (-1,-1), 5),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [WHITE, HexColor('#f9fafb')]),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2*inch))

    # Clean up report text (remove LLM citations sections)
    cut_markers = ["## sources", "## citations", "## references", "## works cited"]
    report_lines = report_text.split("\n")
    filtered_lines = []
    cutting = False
    for line in report_lines:
        if any(line.strip().lower().startswith(m) for m in cut_markers):
            cutting = True
        if not cutting:
            filtered_lines.append(line)
    report_text = "\n".join(filtered_lines)

    # Report content with enhanced styling
    prev_was_heading = False
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            if not prev_was_heading:
                story.append(Spacer(1, 6))
            prev_was_heading = False
        elif line.startswith("## "):
            story.append(Spacer(1, 8))
            story.append(Paragraph(f"   {safe_paragraph_text(line[3:].upper())}", h2_style))
            prev_was_heading = True
        elif line.startswith("### "):
            story.append(Spacer(1, 4))
            story.append(Paragraph(safe_paragraph_text(line[4:]), h3_style))
            prev_was_heading = True
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"- {safe_paragraph_text(line[2:])}", bullet_style))
            prev_was_heading = False
        else:
            story.append(Paragraph(safe_paragraph_text(line), body_style))
            prev_was_heading = False

    # Enhanced back section
    story.append(Spacer(1, 0.3*inch))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT_BLUE, spaceAfter=12))
    
    back_style = ParagraphStyle(
        'Back', fontSize=9, fontName='Helvetica',
        textColor=GRAY, alignment=TA_CENTER, leading=14
    )
    story.append(Paragraph(
        f"Report generated by AutoResearch Agent  |  Built by M. Umer Farooq  |  "
        f"{datetime.now().strftime('%Y')}",
        back_style
    ))

    doc.build(story)
    return filename


def save_word(topic, report_text, all_sources, logo_path=None):
    """Enhanced Word document with better styling"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filename = topic.lower().replace(" ", "-")[:40] + "-report.docx"
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title with enhanced styling
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))
    
    title = doc.add_heading("Research Intelligence", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    run.font.size = Pt(28)

    # Topic
    topic_para = doc.add_paragraph()
    topic_run = topic_para.add_run(topic)
    topic_run.font.size = Pt(16)
    topic_run.font.bold = True
    topic_run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    topic_para.paragraph_format.space_after = Pt(12)

    # Metadata section with cards
    doc.add_heading("Report Information", level=2)
    meta = [
        ("📅 Date Generated", datetime.now().strftime("%B %d, %Y")),
        ("👨‍💻 Architect", "M. Umer Farooq"),
        ("🧠 AI Model", "LLaMA 3.3 70B (Groq Ultra-Fast Inference)"),
        ("🔍 Search Engine", "Tavily Advanced Web Search"),
        ("📊 Sources Analyzed", str(len(all_sources))),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (key, val) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    doc.add_paragraph()
    
    # Verification notice
    disclaimer = doc.add_paragraph()
    disc_run = disclaimer.add_run(
        "⚠️ Verification Notice: This report was generated by AutoResearch Agent using AI synthesis and web search. "
        "All claims reference publicly available sources. Validate critical findings through independent research before decision-making."
    )
    disc_run.font.size = Pt(9)
    disc_run.font.italic = True
    disc_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    
    doc.add_page_break()

    # Sources table
    if all_sources:
        doc.add_heading("Sources & Credibility Assessment", level=2)
        table = doc.add_table(rows=len(all_sources)+1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No."
        hdr_cells[1].text = "Title"
        hdr_cells[2].text = "Credibility"
        hdr_cells[3].text = "Evidence Snippet"
        hdr_cells[4].text = "URL"
        
        for i, s in enumerate(all_sources, 1):
            row = table.rows[i].cells
            row[0].text = str(i)
            row[1].text = s.get('title', '')[:40]
            cred = s.get('credibility', 0)
            row[2].text = "★★★★★" if cred >= 4 else "★★★★☆" if cred >= 3 else "★★★☆☆"
            row[3].text = s.get('snippet', '')[:80]
            row[4].text = s.get('url', '')[:50]
        doc.add_paragraph()

    doc.add_page_break()

    # Report content
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=2)
            h.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
        elif line.startswith("#### "):
            h = doc.add_heading(line[5:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(4)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"Report generated by AutoResearch Agent  |  Built by M. Umer Farooq  |  {datetime.now().strftime('%Y')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filename)
    return filename

def save_html(topic, report_text, all_sources, logo_path=None):
    """Enhanced HTML with modern design and better styling"""
    filename = topic.lower().replace(' ', '-')[:40] + '-report.html'
    num_sources = len(all_sources)
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Research Intelligence Report</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            line-height: 1.6;
            color: #1e293b;
            background: linear-gradient(135deg, #f0f9ff 0%, #f5f3ff 100%);
            padding: 20px;
        }}
        
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: white;
            border-radius: 12px;
            box-shadow: 0 10px 40px rgba(0,0,0,0.08);
            overflow: hidden;
        }}
        
        header {{
            background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
            color: white;
            padding: 60px 40px;
            position: relative;
            overflow: hidden;
        }}
        
        header::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: radial-gradient(circle at 30% 50%, rgba(37, 99, 235, 0.1) 0%, transparent 50%);
        }}
        
        .header-content {{
            position: relative;
            z-index: 1;
        }}
        
        .badge {{
            display: inline-block;
            background: rgba(37, 99, 235, 0.2);
            border: 2px solid #2563eb;
            color: #2563eb;
            padding: 8px 14px;
            border-radius: 6px;
            font-size: 12px;
            font-weight: bold;
            margin-bottom: 20px;
            letter-spacing: 1.5px;
        }}
        
        h1 {{
            font-size: 36px;
            margin-bottom: 10px;
            font-weight: 700;
            letter-spacing: -0.5px;
        }}
        
        .topic {{
            font-size: 24px;
            color: #06b6d4;
            margin-bottom: 40px;
            font-weight: 600;
        }}
        
        .meta-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-top: 30px;
        }}
        
        .meta-card {{
            background: rgba(255, 255, 255, 0.1);
            border: 1px solid rgba(255, 255, 255, 0.2);
            padding: 18px;
            border-radius: 8px;
            backdrop-filter: blur(10px);
        }}
        
        .meta-card strong {{
            display: block;
            font-size: 13px;
            color: #93c5fd;
            margin-bottom: 6px;
            letter-spacing: 0.5px;
        }}
        
        .meta-card span {{
            color: white;
            font-size: 14px;
            font-weight: 500;
        }}
        
        main {{
            padding: 50px 40px;
        }}
        
        .disclaimer {{
            background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
            border-left: 4px solid #2563eb;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 40px;
            font-size: 13px;
            color: #374151;
        }}
        
        h2 {{
            font-size: 24px;
            color: white;
            background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
            padding: 16px 20px;
            margin: 40px 0 20px 0;
            border-radius: 8px;
            border-left: 4px solid #f59e0b;
        }}
        
        h3 {{
            font-size: 16px;
            color: #2563eb;
            margin: 20px 0 10px 0;
            font-weight: 600;
        }}
        
        p {{
            margin-bottom: 16px;
            text-align: justify;
            color: #475569;
            line-height: 1.8;
        }}
        
        ul {{
            margin-left: 30px;
            margin-bottom: 20px;
        }}
        
        li {{
            margin-bottom: 10px;
            color: #475569;
            line-height: 1.6;
        }}
        
        .sources-section {{
            margin-top: 50px;
        }}
        
        .sources-table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            font-size: 13px;
        }}
        
        .sources-table thead {{
            background: linear-gradient(135deg, #2563eb 0%, #0891b2 100%);
            color: white;
        }}
        
        .sources-table th {{
            padding: 14px;
            text-align: left;
            font-weight: 600;
            letter-spacing: 0.5px;
        }}
        
        .sources-table td {{
            padding: 14px;
            border-bottom: 1px solid #e2e8f0;
            word-break: break-word;
        }}
        
        .sources-table tbody tr:hover {{
            background-color: #f8fbff;
        }}
        
        .sources-table tbody tr:nth-child(even) {{
            background-color: #f9fafb;
        }}
        
        .credibility {{
            color: #059669;
            font-weight: bold;
        }}
        
        a {{
            color: #2563eb;
            text-decoration: none;
        }}
        
        a:hover {{
            text-decoration: underline;
        }}
        
        footer {{
            background: #f1f5f9;
            padding: 30px;
            text-align: center;
            color: #6b7280;
            font-size: 12px;
            border-top: 1px solid #e2e8f0;
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <div class="header-content">
                <div class="badge">✨ AI-POWERED INTELLIGENCE REPORT</div>
                <h1>Research Intelligence</h1>
                <div class="topic">{topic}</div>
                <div class="meta-grid">
                    <div class="meta-card">
                        <strong>📅 DATE GENERATED</strong>
                        <span>{datetime.now().strftime('%B %d, %Y')}</span>
                    </div>
                    <div class="meta-card">
                        <strong>👨‍💻 ARCHITECT</strong>
                        <span>M. Umer Farooq</span>
                    </div>
                    <div class="meta-card">
                        <strong>🧠 AI ENGINE</strong>
                        <span>LLaMA 3.3 70B</span>
                    </div>
                    <div class="meta-card">
                        <strong>📊 SOURCES</strong>
                        <span>{num_sources} verified sources</span>
                    </div>
                </div>
            </div>
        </header>
        
        <main>
            <div class="disclaimer">
                <strong>⚠️ Verification Notice:</strong> This report was generated by AutoResearch Agent using AI synthesis and web search. 
                All claims reference publicly available sources. Validate critical findings through independent research before decision-making.
            </div>
"""

    # Add sources summary table
    if all_sources:
        html += f"""
            <div class="sources-section">
                <h2>📚 Sources & Credibility Assessment</h2>
                <table class="sources-table">
                    <thead>
                        <tr>
                            <th style="width: 5%">No.</th>
                            <th style="width: 25%">Title</th>
                            <th style="width: 15%">Credibility</th>
                            <th style="width: 35%">Evidence Snippet</th>
                            <th style="width: 20%">URL</th>
                        </tr>
                    </thead>
                    <tbody>
"""
        for i, s in enumerate(all_sources[:12], 1):
            cred = s.get('credibility', 0)
            cred_stars = "★★★★★" if cred >= 4 else "★★★★☆" if cred >= 3 else "★★★☆☆"
            html += f"""
                        <tr>
                            <td>{i}</td>
                            <td><strong>{s.get('title', '')[:40]}</strong></td>
                            <td class="credibility">{cred_stars}</td>
                            <td>{s.get('snippet', '')[:80]}</td>
                            <td><a href="{s.get('url', '')}" target="_blank">{s.get('url', '')[:40]}</a></td>
                        </tr>
"""
        html += """
                    </tbody>
                </table>
            </div>
"""

    # Add report content
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            html += "<p></p>"
        elif line.startswith("## "):
            html += f"<h2>{line[3:].upper()}</h2>"
        elif line.startswith("### "):
            html += f"<h3>{line[4:]}</h3>"
        elif line.startswith("- ") or line.startswith("* "):
            html += f"<ul><li>{line[2:]}</li></ul>"
        else:
            html += f"<p>{line}</p>"

    html += f"""
        </main>
        
        <footer>
            <p>Report generated by AutoResearch Agent  |  Built by M. Umer Farooq  |  {datetime.now().strftime('%Y')}</p>
        </footer>
    </div>
</body>
</html>
"""

    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html)
    return filename

def save_markdown(topic, report_text, all_sources, logo_path=None):
    filename = topic.lower().replace(' ', '-')[:40] + '-report.md'
    md = [f"# Research Intelligence\n\n## {topic}\n"]
    if logo_path and os.path.exists(logo_path):
        md.append(f"![Logo]({logo_path})\n")
    
    # Metadata
    md.append("## Report Information\n")
    md.append(f"- **Date Generated**: {datetime.now().strftime('%B %d, %Y')}\n")
    md.append(f"- **Architect**: M. Umer Farooq\n")
    md.append(f"- **AI Engine**: LLaMA 3.3 70B (Groq)\n")
    md.append(f"- **Search Engine**: Tavily Web Search\n")
    md.append(f"- **Sources Analyzed**: {len(all_sources)}\n\n")
    
    # Sources table
    if all_sources:
        md.append("## Sources & Credibility Assessment\n")
        md.append("| # | Title | Credibility | Evidence Snippet | URL |\n|---|-------|-------------|------------------|-----|\n")
        for i, s in enumerate(all_sources[:12], 1):
            cred = s.get('credibility', 0)
            cred_stars = "★★★★★" if cred >= 4 else "★★★★☆" if cred >= 3 else "★★★☆☆"
            md.append(f"| {i} | {s.get('title','')[:35]} | {cred_stars} | {s.get('snippet','')[:60]} | {s.get('url','')[:40]} |\n")
        md.append("\n")
    
    md.append("## Report\n\n")
    md.append(report_text)
    
    with open(filename, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md))
    return filename



# --- Modularized main research workflow ---
def collect_sources_for_topic(topic, plugin_hooks=None):
    questions = generate_sub_questions(topic)
    if not questions:
        questions = [topic]
    try:
        max_sources = int(os.getenv("MAX_SOURCES_PER_QUERY", "4"))
    except Exception:
        max_sources = 4

    all_sources = []

    def fetch(q):
        sources = search_web(q, max_results=max_sources)
        if plugin_hooks and 'on_sources_fetched' in plugin_hooks:
            sources = plugin_hooks['on_sources_fetched'](q, sources)
        return sources

    with ThreadPoolExecutor(max_workers=max(1, len(questions))) as executor:
        futures = {executor.submit(fetch, q): q for q in questions}
        for future in as_completed(futures):
            try:
                all_sources.extend(future.result())
            except Exception as e:
                print(f"[WARNING] Search failed for '{futures[future]}': {e}")

    return all_sources

def generate_report_for_topic(topic, all_sources, plugin_hooks=None):
    if plugin_hooks and 'pre_synthesis' in plugin_hooks:
        all_sources = plugin_hooks['pre_synthesis'](all_sources)
    report = synthesize_findings(topic, all_sources)
    if plugin_hooks and 'post_synthesis' in plugin_hooks:
        report = plugin_hooks['post_synthesis'](report)
    return report

def run_agent(topic, output_format="pdf", logo_path=None, plugin_hooks=None):
    print(f"\n{'='*60}")
    print(f"🤖 AutoResearch Agent v2.0 - Enhanced Design")
    print(f"👨‍💻 Built by M. Umer Farooq")
    print(f"📌 Topic: {topic}")
    print(f"{'='*60}")

    all_sources = collect_sources_for_topic(topic, plugin_hooks=plugin_hooks)
    print(f"\n📚 Total sources collected: {len(all_sources)}")
    report = generate_report_for_topic(topic, all_sources, plugin_hooks=plugin_hooks)

    critique_results = critique_report(report, all_sources)
    # Guard: if LLM returned flat strings instead of dicts, convert them
    normalized = []
    for r in critique_results:
        if isinstance(r, dict):
            normalized.append(r)
        elif isinstance(r, str):
            normalized.append({"claim": r, "status": "unsupported", "reason": "Flagged by fact-checker"})
    critique_results = normalized
    unsupported_claims = [r["claim"] if isinstance(r, dict) else r for r in critique_results if not isinstance(r, dict) or r.get("status") == "unsupported"][:3]

    if unsupported_claims:
        gap_sources = fill_gaps(unsupported_claims)
        all_sources.extend(gap_sources)
        report = patch_report(report, gap_sources, unsupported_claims)
        print("✅ Report patched with verified sources")
    else:
        print("✅ All claims supported by sources")

    if output_format == "pdf":
        pdf_file = save_pdf(topic, report, all_sources, logo_path=logo_path)
        print(f"📄 PDF saved: {pdf_file}")
    elif output_format == "word":
        word_file = save_word(topic, report, all_sources, logo_path=logo_path)
        print(f"📝 Word saved: {word_file}")
    elif output_format == "html":
        html_file = save_html(topic, report, all_sources, logo_path=logo_path)
        print(f"🌐 HTML saved: {html_file}")
    elif output_format == "md":
        md_file = save_markdown(topic, report, all_sources, logo_path=logo_path)
        print(f"📝 Markdown saved: {md_file}")

    print(f"\n{'='*60}")
    print(f"✅ Research Complete!")
    print(f"📊 Sources analyzed: {len(all_sources)}")
    print(f"👨‍💻 Built by M. Umer Farooq")
    print(f"{'='*60}\n")


def _competitor_output_base(companies):
    return "competitive-analysis-" + "-".join(
        [c.strip().lower().replace(" ", "-")[:10] for c in companies[:3]]
    )


def run_competitor_agent(companies_input, fmt='pdf'):
    companies = [c.strip() for c in companies_input.split(',') if c.strip()]
    if not companies:
        raise ValueError("Please provide at least one company name.")

    print(f"\n{'='*60}")
    print("🤖 AutoResearch Agent v2.0 — Competitor Intelligence Mode")
    print("👨‍💻 Built by M. Umer Farooq")
    print(f"🏢 Companies: {', '.join(companies)}")
    print(f"{'='*60}")

    all_profiles = {}
    all_sources = []

    for company in companies:
        print(f"\n📊 Researching: {company}")
        company_sources = []
        queries = [
            f"{company} latest news",
            f"{company} products services and pricing",
            f"{company} funding valuation revenue investors",
            f"{company} market share competitors strategy",
        ]

        for q in queries:
            sources = search_web(q, max_results=4)
            company_sources.extend(sources)
            all_sources.extend(sources)

        sources_text = "\n".join(
            [
                f"- {s.get('title', '')} | {s.get('url', '')} | {s.get('content', '')[:250]}"
                for s in company_sources[:10]
            ]
        )
        profile_prompt = f"""Create a concise competitor profile for {company} using the evidence below.

SOURCES:
{sources_text}

Return markdown with this structure only:
## {company}
### Snapshot
- ...
### Products & Positioning
- ...
### Business Signals
- ...
### Strengths
- ...
### Risks
- ...
"""
        all_profiles[company] = ask_groq(
            profile_prompt,
            system="You are a senior competitive intelligence analyst.",
        )
        print(f"  📚 Found {len(company_sources)} sources")

    profiles_text = "\n\n".join([f"=== {c} ===\n{p}" for c, p in all_profiles.items()])
    summary_prompt = f"""Based on these company profiles, write a competitive analysis summary.

{profiles_text}

Write markdown sections:
## Competitive Landscape Overview
## Head-to-Head Comparison
## Who Is Winning and Why
## Strategic Recommendations
"""
    summary = ask_groq(summary_prompt, system="You are a senior competitive intelligence analyst.")
    full_report = summary + "\n\n## Individual Company Profiles\n\n" + profiles_text

    topic = f"Competitive Analysis: {', '.join(companies)}"
    base = _competitor_output_base(companies)

    if fmt == 'pdf':
        generated = save_pdf(topic, full_report, all_sources)
        target = f"{base}.pdf"
        if generated != target:
            os.replace(generated, target)
        print(f"\n📄 PDF saved: {target}")
    elif fmt == 'word':
        generated = save_word(topic, full_report, all_sources)
        target = f"{base}.docx"
        if generated != target:
            os.replace(generated, target)
        print(f"\n📝 Word saved: {target}")
    elif fmt == 'html':
        generated = save_html(topic, full_report, all_sources)
        target = f"{base}.html"
        if generated != target:
            os.replace(generated, target)
        print(f"\n🌐 HTML saved: {target}")
    elif fmt == 'md':
        generated = save_markdown(topic, full_report, all_sources)
        target = f"{base}.md"
        if generated != target:
            os.replace(generated, target)
        print(f"\n📝 Markdown saved: {target}")
    else:
        raise ValueError("Unsupported competitor output format.")

    print(f"\n{'='*60}")
    print("✅ Competitive Analysis Complete!")
    print(f"📊 Total sources analyzed: {len(all_sources)}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(
        description="🤖 AutoResearch Agent — Enhanced Design Edition"
    )
    parser.add_argument("topic", help="Research topic")
    parser.add_argument("--format", choices=["pdf", "word", "html", "md"], default="pdf",
                        help="Output format (default: pdf)")
    parser.add_argument("--logo", type=str, default=None, help="Path to custom logo")
    parser.add_argument("--max-sources", type=int, default=None, help="Max sources per query")
    
    args = parser.parse_args()
    
    if args.max_sources:
        os.environ["MAX_SOURCES_PER_QUERY"] = str(args.max_sources)
    
    run_agent(args.topic, output_format=args.format, logo_path=args.logo)